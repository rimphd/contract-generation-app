#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import requests
from io import BytesIO
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from dotenv import load_dotenv
import json, os
from pathlib import Path

load_dotenv()
MODELS_JSON_PATH = os.getenv("MODELS_JSON", "models.json")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
def load_models_from_json(path: str | Path):
    path = Path(path)
    default_model = "meta-llama/llama-3.1-70b-instruct"
    models_list = []

    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        default_model = data.get("default", default_model)
        for item in data.get("models", []):
            mid = item.get("id")
            label = item.get("label", mid)
            if mid:
                models_list.append((mid, label))
    except Exception as e:
        # En cas d'erreur, on garde un fallback minimal
        print(f"[WARN] models.json non lu ({e}); utilisation d’un fallback.")
        models_list = [(default_model, "Llama 3.1 70B Instruct (Meta)")]

    # Si le default n'est pas dans la liste, on l’ajoute en tête
    if models_list and default_model not in {m[0] for m in models_list}:
        models_list.insert(0, (default_model, f"{default_model}"))

    return default_model, models_list
DEFAULT_MODEL, AVAILABLE_MODELS = load_models_from_json(MODELS_JSON_PATH)
app = Flask(__name__)
app.secret_key = "dev-secret"  # change en prod

def call_openrouter(model_id: str, prompt: str, temperature=0.4, max_tokens=1600) -> str:
    headers = {
        "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY')}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost",
        "X-Title": "Contract-UI",
    }
    payload = {
        "model": model_id,
        "messages": [
            {"role": "system", "content": "Tu es un assistant juridique et tu rédiges des contrats complets et soignés."},
            {"role": "user", "content": prompt},
        ],
        "temperature": float(temperature),   # <- déjà normalisé
        "max_tokens": int(max_tokens),
        "stream": False,
    }

    r = requests.post("https://openrouter.ai/api/v1/chat/completions",
                      headers=headers, json=payload, timeout=180)
    if r.status_code != 200:
        # Affiche l'erreur JSON d’OpenRouter pour diagnostiquer (évite le 400 silencieux)
        try:
            err = r.json()
        except Exception:
            err = r.text
        raise RuntimeError(f"OpenRouter {r.status_code}: {err}")

    data = r.json()
    return data["choices"][0]["message"]["content"].strip()

def build_prompt(p):
    return f"""
Tu es un assistant juridique. Génère un contrat de location (contrat de bail) clair et professionnel.

Paramètres:
- LOCATAIRE: {p['tenant_name']}
- BAILLEUR: {p['landlord_name']}
- LOYER_MENSUEL: {p['rent']} MAD
- DEPOT_DE_GARANTIE: {p['security_deposit']} MAD
- DUREE_MOIS: {p['duration_months']} mois
- ADRESSE: {p['address']}
- DATE_DEBUT: {p['start_date']}

Exigences de sortie:
- Rédige le contrat COMPLET en français, style formel (1–2 pages).
- Pas de JSON ni de code block. Retourne du TEXTE pur prêt à copier.
- Inclure: identité des parties, objet/adresse du bien, durée/renouvellement, loyer et paiement,
  dépôt de garantie, obligations bailleur/locataire, réparations/charges, résiliation/préavis,
  état des lieux, clause de juridiction/applicable, signatures (placeholders).
""".strip()

def make_docx(title: str, text: str) -> BytesIO:
    buf = BytesIO()
    doc = Document()
    doc.add_heading(title, level=0)
    for para in text.split("\n\n"):
        doc.add_paragraph(para.strip())
    doc.save(buf)
    buf.seek(0)
    return buf

def make_pdf(title: str, text: str) -> BytesIO:
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.5*cm)]
    for para in text.split("\n\n"):
        story.append(Paragraph(para.replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 0.3*cm))
    doc.build(story)
    buf.seek(0)
    return buf
def to_float_fr(x, default=None):
    """Accepte '0,4' '0.4' ' 0,4 ' -> 0.4"""
    try:
        if x is None:
            return default
        s = str(x).strip().replace("\u00a0", "").replace(" ", "").replace(",", ".")
        return float(s)
    except Exception:
        return default

def to_int_fr(x, default=None):
    """Accepte '7 000' '7,000' '7000' -> 7000"""
    try:
        if x is None:
            return default
        s = str(x).strip().replace("\u00a0", "").replace(" ", "").replace(",", "")
        return int(s)
    except Exception:
        return default
@app.route("/", methods=["GET"])
def index():
    return render_template("index.html",
                           default_model=DEFAULT_MODEL,
                           models=AVAILABLE_MODELS)


@app.route("/generate", methods=["POST"])
def generate():
    form = request.form

    params = {
        "tenant_name": form.get("tenant_name","").strip(),
        "landlord_name": form.get("landlord_name","").strip(),
        "rent": to_int_fr(form.get("rent")),
        "security_deposit": to_int_fr(form.get("security_deposit")),
        "duration_months": to_int_fr(form.get("duration_months")),
        "address": form.get("address","").strip(),
        "start_date": form.get("start_date","").strip(),
    }

    # validation
    missing = [k for k in ["tenant_name","landlord_name","address","start_date"] if not params[k]]
    invalid = [k for k in ["rent","security_deposit","duration_months"] if params[k] is None]
    if missing or invalid:
        msg = []
        if missing: msg.append(f"Champs manquants : {', '.join(missing)}")
        if invalid: msg.append(f"Champs numériques invalides : {', '.join(invalid)}")
        flash(" — ".join(msg), "danger")
        return redirect(url_for("index"))

    # modèle + température (temp peut être '0,4' -> 0.4)
    model_choice = (request.form.get("model_id") or "").strip()
    if model_choice == "__custom__":
        model_id = (request.form.get("model_id_custom") or "").strip()
    else:
        model_id = model_choice or DEFAULT_MODEL
    temperature = to_float_fr(form.get("temperature"), 0.4)

    prompt = build_prompt(params)

    try:
        contract_text = call_openrouter(model_id, prompt, temperature=temperature, max_tokens=1600)
    except Exception as e:
        flash(str(e), "danger")
        return redirect(url_for("index"))

    return render_template("result.html",
                           contract_text=contract_text,
                           params=params,
                           model_id=model_id,
                           temperature=temperature)


@app.post("/download-docx")
def download_docx_no_db():
    text = request.form.get("contract_text", "")
    if not text.strip():
        flash("Pas de contrat à télécharger.", "danger")
        return redirect(url_for("index"))
    buf = make_docx("Contrat de location", text)
    return send_file(buf,
                     as_attachment=True,
                     download_name="contrat.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/download-pdf")
def download_pdf_no_db():
    text = request.form.get("contract_text", "")
    if not text.strip():
        flash("Pas de contrat à télécharger.", "danger")
        return redirect(url_for("index"))
    buf = make_pdf("Contrat de location", text)
    return send_file(buf,
                     as_attachment=True,
                     download_name="contrat.pdf",
                     mimetype="application/pdf")

if __name__ == "__main__":
    app.run(debug=True)
